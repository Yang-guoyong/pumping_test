"""
拐点法求参程序

MIT License

版权所有 (c) 2025 yanggy1010 (yanggy1010@cumt.edu.cn)

特此免费授予任何获得本软件及相关文档文件（以下简称“软件”）副本的人，不受限制地处理本软件的权限，包括
但不限于使用、复制、修改、合并、发布、分发、再许可和/或出售本软件的副本，并允许接收本软件的人这样做，
但须符合以下条件：

上述版权声明和本许可声明应包含在本软件的所有副本或实质性部分中。

本软件按“原样”提供，不提供任何形式的明示或暗示担保，包括但不限于适销性、特定用途适用性和非侵权性的担保。
在任何情况下，作者或版权持有人均不对任何索赔、损害或其他责任负责，无论是在合同、侵权或其他行为中产生的，
还是与本软件或本软件的使用或其他交易有关的。
"""

import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
from itertools import zip_longest

from scipy.special import expn, k0
from scipy.optimize import bisect

# 配置字体和全局参数
py_path = os.path.dirname(__file__)
font_mapping = {
    "simkai.ttf": "KaiTi", 
    "simhei.ttf": "simhei",         
    "simfang.ttf": "FangSong"  
}
loaded_font_families = ['Times New Roman']
for font_file, font_family in font_mapping.items():
    font_path = os.path.join(py_path, "fonts", font_file)
    if os.path.exists(font_path):
        font_manager.fontManager.addfont(font_path)
        loaded_font_families.append(font_family)
        # break
plt.rcParams['font.family'] = loaded_font_families
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['xtick.direction'] = 'in'
plt.rcParams['ytick.direction'] = 'in'

# Hantush-Jacob 井函数计算
def hantush_jacob_well(u, beta):
    """计算 Hantush-Jacob 井函数。"""
    def wellfunc(u, beta):
        if u < 0:
            raise ValueError("Negative values for 'u' are not allowed.")
        if u == 0:
            return 2.0 * k0(beta)
        r, b, t = 1.0, 2 * u, beta**2 / (4 * u)
        W, n = 0.0, 0
        if beta <= b:
            term = r * expn(n + 1, u)
            while np.abs(term) > 1e-10:
                W += term
                n += 1
                r *= -t / n
                term = r * expn(n + 1, u)
        else:
            W = 2.0 * k0(beta)
            term = r * expn(n + 1, t)
            while np.abs(term) > 1e-10:
                W -= term
                n += 1
                r *= -u / n
                term = r * expn(n + 1, t)
        return W
    return np.vectorize(wellfunc)(u, beta)

def calc_init_params(t, s):
    """按 Jacob 公式计算参数。"""
    sp = 0.5 * s[-1]
    idx = np.where(s > sp)[0][0] - 1
    tp = t[idx]
    slope = (s[idx] - s[idx + 1]) / np.log10(t[idx] / t[idx + 1])
    return sp, tp, slope

# 生成Word文档
def create_report(T, S, B, fig, filename="report.docx"):
    """生成Word文档"""
    doc = Document()
    doc.add_heading("配线法求参", level=1)
    doc.add_paragraph(f"导水系数 T = {T:.4e} m²/min")
    doc.add_paragraph(f"贮水系数 S = {S:.4e}")
    doc.add_paragraph(f"越流因素 B = {B:.4e} m")

    img_buffer = BytesIO()
    fig.savefig(img_buffer, dpi=300, format="png", bbox_inches="tight")
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    
    report_buffer = BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)
    return report_buffer

# 主程序
def inflection():
    st.markdown("### 拐点法求参")

    t = np.array([1, 4, 7, 10, 15, 20, 25, 30, 45, 60, 
                  75, 90, 120, 150, 180, 210, 240, 270, 300, 330,
                  360, 390, 420, 450, 480, 510, 540])
    s = np.array([0.05, 0.054, 0.12, 0.175, 0.26, 0.33, 0.383, 0.425, 0.52, 
                  0.575, 0.62, 0.64, 0.685, 0.725, 0.735, 0.755, 0.76, 0.76,
                  0.763, 0.77, 0.772, 0.785, 0.79, 0.792, 0.794, 0.795, 0.796])
    Q = 69.1 / 60  # m³/min
    r = 197.0      # m

    data = [t, s, [Q], [r]]
    transposed_data = list(map(list, zip_longest(*data, fillvalue=None)))
    df = pd.DataFrame(transposed_data, columns=['t', 's', 'Q', 'r'])

    # 上传数据
    uploaded_file = st.sidebar.file_uploader(
        "上传 CSV 文件", 
        type=["csv"])

    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        t = df['t'].values   # 时间
        s = df['s'].values   # 降深
        Q = df['Q'].iloc[0]  # 抽水量
        r = df['r'].iloc[0]  # 观测距离
    
    st.sidebar.write("CSV数据预览:")
    st.sidebar.dataframe(df, height=200, width=400, hide_index=True) 

    sp, tp, slope = calc_init_params(t, s)

    tp = st.sidebar.slider(r'$t_p$', float(tp/3), float(tp*3), float(tp), float(tp/50), format="%.3f")
    sp = st.sidebar.number_input(r'$s_p$', float(sp), float(sp*2), float(sp), format="%.3f")
    slope = st.sidebar.slider(r'$i$', float(slope/3), float(slope*3), float(slope), float(slope/50), format="%.3f")

    help_txt = '''
        - 通过滑块调整拐点与切线斜率;
        - 拐点降深需要输入;
        - 数据为 csv 格式, 需包含 4 列数据: 时间(min), 降深(m), 抽水量(m³/min), 观测距离(m); 
        可下载预览中的数据来了解该文件格式。 
        '''
    with st.sidebar:
        st.markdown(help_txt)

    # 计算和绘图
    beta = bisect(lambda x: np.exp(x) * k0(x) - 2.3 * sp / slope, 0.001, 5)
    T = 2.3 * Q * np.exp(-beta) / (4 * np.pi * slope)
    S = 2 * T * tp * beta / r**2
    B = r / beta

    fig, ax = plt.subplots(figsize=(6, 4), dpi=150)
    ax.plot(t, s, '*', label="观测值")
    ax.plot(tp, sp, 'o', label="拐点")
    x_log_min = np.floor(np.log10(min(t)))
    x_log_max = np.ceil(np.log10(max(t)))
    t_ = 10**np.linspace(x_log_min, x_log_max, 100)
    ax.plot(t_, sp + slope * np.log10(t_ / tp), label="拐点切线")
    u = 0.25 * r**2 * S / (T * t_)
    ax.plot(t_, 0.25 * Q * hantush_jacob_well(u, beta) / (np.pi * T), 
            label="标准曲线", linestyle="-")
    
    ax.set(xscale="log", 
           xlim=(10**x_log_min, 10**x_log_max), 
           ylim=(0, np.ceil(max(s) * 10) / 10),
                  xlabel=r'$\lg t$',
           ylabel=r'$s$'
    )

    ax.grid(True, which="major", linestyle="-", linewidth=0.5)
    ax.grid(True, which="minor", linestyle="-", linewidth=0.2)
    ax.legend(loc=4)
    
    st.pyplot(fig)
    st.write(f'T = {T:.4f} m²/min, S = {S:.4e}, B = {B:.4e} m')

    # 生成Word文档
    if st.button("Word文档"):
        report = create_report(T, S, B, fig)
        st.download_button(
            label="下载文档",
            data=report,
            file_name="report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# 运行程序
if __name__ == "__main__":
    inflection()