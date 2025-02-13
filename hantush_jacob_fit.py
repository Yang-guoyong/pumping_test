"""
Theis 与 Hantush-Jacob 配线法求参程序

MIT License

版权所有 (c) 2025 yanggy1010 (yanggy1010@cumt.edu.cn)

特此免费授予任何获得本软件及相关文档文件（以下简称“软件”）副本的人，不受限制地处理本软件的权限，包括
但不限于使用、复制、修改、合并、发布、分发、再许可和/或出售本软件的副本，并允许接收本软件的人这样做，
但须符合以下条件：

上述版权声明和本许可声明应包含在本软件的所有副本或实质性部分中。

本软件按“原样”提供，不提供任何形式的明示或暗示担保，包括但不限于适销性、特定用途适用性和非侵权性的担保。
在任何情况下，作者或版权持有人均不对任何索赔、损害或其他责任负责，无论是在合同、侵权或其他行为中产生的，
还是与本软件或本软件的使用或其他交易有关的。
"""

import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
from itertools import zip_longest

from scipy.special import expn, k0, exp1 as theis_well

# 配置字体和全局参数
py_path = os.path.dirname(__file__)
font_mapping = {
    "simkai.ttf": "KaiTi", 
    "simhei.ttf": "simhei",         
    "simfang.ttf": "FangSong"  
}
loaded_font_families = ['Times New Roman']
for font_file, font_family in font_mapping.items():
    font_path = os.path.join(py_path, "fonts", font_file)
    if os.path.exists(font_path):
        font_manager.fontManager.addfont(font_path)
        loaded_font_families.append(font_family)
        # break
plt.rcParams['font.family'] = loaded_font_families
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['xtick.direction'] = 'in'
plt.rcParams['ytick.direction'] = 'in'

# Hantush-Jacob 井函数计算
def hantush_jacob_well(u, beta):
    """计算 Hantush-Jacob 井函数。"""
    def wellfunc(u, beta):
        if u < 0:
            raise ValueError("Negative values for 'u' are not allowed.")
        if u == 0:
            return 2.0 * k0(beta)
        r, b, t = 1.0, 2 * u, beta**2 / (4 * u)
        W, n = 0.0, 0
        if beta <= b:
            term = r * expn(n + 1, u)
            while np.abs(term) > 1e-10:
                W += term
                n += 1
                r *= -t / n
                term = r * expn(n + 1, u)
        else:
            W = 2.0 * k0(beta)
            term = r * expn(n + 1, t)
            while np.abs(term) > 1e-10:
                W -= term
                n += 1
                r *= -u / n
                term = r * expn(n + 1, t)
        return W
    return np.vectorize(wellfunc)(u, beta)

def calc_init_params(t, s, Q, r):
    """按 Jacob 公式计算参数。"""
    idx = np.argmin(np.abs(t / t[-1] - 0.5)) - 1
    slope = (s[idx] - s[idx + 1]) / np.log10(t[idx] / t[idx + 1])
    T = 0.183 * Q / slope
    t0 = t[idx] * 10**(-s[idx] / slope)
    S = 2.25 * T * t0 / r**2
    dx = np.log10(4 * T / (S * r**2))
    dy = np.log10(4 * np.pi * T / Q)
    return T, S, dx, dy

# 生成Word文档
def create_report(T, S, B, fig, filename="report.docx"):
    """生成Word文档"""
    doc = Document()
    doc.add_heading("配线法求参", level=1)
    doc.add_paragraph(f"导水系数 T = {T:.4e} m²/min")
    doc.add_paragraph(f"贮水系数 S = {S:.4e}")
    doc.add_paragraph(f"越流因素 B = {B:.4e} m")

    img_buffer = BytesIO()
    fig.savefig(img_buffer, dpi=300, format="png", bbox_inches="tight")
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    
    report_buffer = BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)
    return report_buffer

# 主程序
def hantush_jacob_fit():
    st.markdown("### Theis 与 Hantush-Jacob 配线法")

    t = np.array([1, 2, 4, 6, 9, 20, 30, 40, 50, 60, 
                  90, 120, 150, 360, 550, 720])  # min
    s = np.array([2.5, 3.9, 6.1, 8.0, 10.6, 16.8, 20.0, 22.6, 24.7, 26.4,
                  30.4, 33.0, 35.0, 42.6, 44.0, 44.5]) / 100   # m
    Q = 0.3667  # 528/1440, m³/min
    r = 90      # m

    data = [t, s, [Q], [r]]
    transposed_data = list(map(list, zip_longest(*data, fillvalue=None)))
    df = pd.DataFrame(transposed_data, columns=['t', 's', 'Q', 'r'])

    # 上传数据
    uploaded_file = st.sidebar.file_uploader(
        "上传 CSV 文件", 
        type=["csv"])

    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        t = df['t'].values   # 时间
        s = df['s'].values   # 降深
        Q = df['Q'].iloc[0]  # 抽水量
        r = df['r'].iloc[0]  # 观测距离
    
    st.sidebar.write("CSV数据预览:")
    st.sidebar.dataframe(df, height=150, width=400, hide_index=True) 

    T, S, dx, dy = calc_init_params(t, s, Q, r)
    dx = st.sidebar.slider(r"$\Delta \lg t$", dx-0.5, dx+0.5, dx, 0.01, format="%.3f")
    dy = st.sidebar.slider(r"$\Delta \lg s$", dy-0.5, dy+0.5, dy, 0.01, format="%.3f")
    beta = st.sidebar.slider(r'$\beta$', 0.0, 2.0, 0.0, step=0.01, format="%.3f")

    help_txt = '''
        - 通过滑块平移图形配线;
        - 数据为 csv 格式, 需包含 4 列数据: 时间(min), 降深(m), 抽水量(m³/min), 观测距离(m); 
        可下载预览中的数据来了解该文件格式。 
        '''
    with st.sidebar:
        st.markdown(help_txt)

    # 计算和绘图
    fig, ax = plt.subplots(dpi=150)
    ax.set(xscale="log", yscale="log", xlim=(1.0e-1, 1.0e4), ylim=(1.0e-2, 10.0),
           xlabel=r'$\lg \frac{1}{u}$', ylabel=r'$\lg W(u)$')
    ax.grid(True, which="major", linestyle="-", linewidth=0.5)
    ax.grid(True, which="minor", linestyle="-", linewidth=0.2)
    ax.set_aspect('equal')

    t_ = 10**np.linspace(-1, 4, 51)
    ax.plot(t_, theis_well(1 / t_), label="Theis", linestyle="--")
    ax.plot(t_, hantush_jacob_well(1 / t_, beta), label="Hantush-Jacob")
    ax.plot(t * 10**dx, s * 10**dy, "r*", label="观测值")
    ax.legend(loc=4)

    st.pyplot(fig)

    # 计算水文参数
    T = Q / (4 * np.pi) * 10**dy
    S = 4 * T / r**2 * 10**(-dx)
    st.write(f'T = {T:.4f} m^2/min, S = {S:.4e}')
    if beta > 0:
        B = r / beta
    else:
        B = float('inf')
    st.write(f'B = {B:.4e}')

    # 生成Word文档
    if st.button("Word文档"):
        report = create_report(T, S, B, fig)
        st.download_button(
            label="下载文档",
            data=report,
            file_name="report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# 运行程序
if __name__ == "__main__":
    hantush_jacob_fit()