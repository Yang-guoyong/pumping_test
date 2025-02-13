"""
Jacob 直线图解法求参程序

MIT License

版权所有 (c) 2025 yanggy1010 (yanggy1010@cumt.edu.cn)

特此免费授予任何获得本软件及相关文档文件（以下简称“软件”）副本的人，不受限制地处理本软件的权限，包括
但不限于使用、复制、修改、合并、发布、分发、再许可和/或出售本软件的副本，并允许接收本软件的人这样做，
但须符合以下条件：

上述版权声明和本许可声明应包含在本软件的所有副本或实质性部分中。

本软件按“原样”提供，不提供任何形式的明示或暗示担保，包括但不限于适销性、特定用途适用性和非侵权性的担保。
在任何情况下，作者或版权持有人均不对任何索赔、损害或其他责任负责，无论是在合同、侵权或其他行为中产生的，
还是与本软件或本软件的使用或其他交易有关的。
"""

import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
from itertools import zip_longest

# 配置字体和全局参数
py_path = os.path.dirname(__file__)
font_mapping = {
    "simkai.ttf": "KaiTi", 
    "simhei.ttf": "simhei",         
    "simfang.ttf": "FangSong"  
}
loaded_font_families = ['Times New Roman']
for font_file, font_family in font_mapping.items():
    font_path = os.path.join(py_path, "fonts", font_file)
    if os.path.exists(font_path):
        font_manager.fontManager.addfont(font_path)
        loaded_font_families.append(font_family)
        # break
plt.rcParams['font.family'] = loaded_font_families
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['xtick.direction'] = 'in'
plt.rcParams['ytick.direction'] = 'in'

# 计算初始参数
def calc_init_param(t, s):
    """按 Jacob 公式计算参数。"""
    idx = np.argmin(np.abs(t / t[-1] - 0.5)) - 1
    slope = (s[idx] - s[idx + 1]) / np.log10(t[idx] / t[idx + 1])
    t0 = t[idx] * 10**(-s[idx] / slope)
    return t0, slope

# 生成Word文档
def create_report(T, S, fig, filename="report.docx"):
    """生成Word文档"""
    doc = Document()
    doc.add_heading("直线图解法求参", level=1)
    doc.add_paragraph(f"导水系数 T = {T:.4e} m²/min")
    doc.add_paragraph(f"贮水系数 S = {S:.4e}")
    
    img_buffer = BytesIO()
    fig.savefig(img_buffer, dpi=300, format="png", bbox_inches="tight")
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    
    report_buffer = BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)
    return report_buffer

# 主程序
def jacob_line():
    st.markdown("### Jacob直线图解法求参")

    t = np.array([1, 2, 4, 6, 9, 20, 30, 40, 50, 60, 
                  90, 120, 150, 360, 550, 720])  # min
    s = np.array([2.5, 3.9, 6.1, 8.0, 10.6, 16.8, 20.0, 22.6, 24.7, 26.4,
                  30.4, 33.0, 35.0, 42.6, 44.0, 44.5]) / 100   # m
    Q = 0.3667  # 528/1440, m³/min
    r = 90      # m

    data = [t, s, [Q], [r]]
    transposed_data = list(map(list, zip_longest(*data, fillvalue=None)))
    df = pd.DataFrame(transposed_data, columns=['t', 's', 'Q', 'r'])

    # 上传数据
    uploaded_file = st.sidebar.file_uploader(
        "上传 CSV 文件", 
        type=["csv"])

    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        t = df['t'].values   # 时间
        s = df['s'].values   # 降深
        Q = df['Q'].iloc[0]  # 抽水量
        r = df['r'].iloc[0]  # 观测距离
    
    st.sidebar.write("CSV数据预览:")
    st.sidebar.dataframe(df, height=200, width=400, hide_index=True) 

    t0, slope = calc_init_param(t, s)
    t0 = st.sidebar.slider(r'$t_0$', t0/3, t0*3, t0, t0/50, format="%.3f")
    slope = st.sidebar.slider(r'$i$', slope/3, slope*3, slope, slope/50, format="%.3f")

    help_txt = '''
        - 通过滑块调整零降深时间与直线斜率;
        - 抽水初始阶段不满足 u 比较小的条件, 这些观测数据不能用;
        - 长时间抽水, 降落漏斗扩展到边界, 后期降深变化平缓的观测数据不能用。
        - 数据为 csv 格式, 需包含 4 列数据: 时间(min), 降深(m), 抽水量(m³/min), 观测距离(m); 
        可下载预览中的数据来了解该文件格式。 
        '''
    with st.sidebar:
        st.markdown(help_txt)

    # 计算和绘图
    T = 0.183 * Q / slope
    S = 2.25 * T * t0 / r**2 

    fig, ax = plt.subplots(figsize=(6, 4), dpi=150)
    ax.set(xscale="log", xlabel=r'$\lg t$', ylabel=r'$s$')
    ax.grid(True, which="major", linestyle="-", linewidth=0.5)
    ax.grid(True, which="minor", linestyle="-", linewidth=0.2)
    ax.plot(t, s, "r*", label="观测值")

    t_ = np.linspace(min(t), max(t), 50)
    s_ = slope * np.log10(t_ / t0)
    ax.plot(t_, s_, label="拟合直线")
    ax.legend(loc=4)

    st.pyplot(fig)
    st.write(f'T = {T:.4f} m²/min, S = {S:.4e}')

    # 生成Word文档
    if st.button("Word文档"):
        report = create_report(T, S, fig)
        st.download_button(
            label="下载文档",
            data=report,
            file_name="report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# 运行程序
if __name__ == "__main__":
    jacob_line()
